"""Word (.docx) export — premium narrative report that follows the WHO-EMRO
'EMR Toolkit SNV narrative template' structure (front matter + chapters 1-9), with an
automatic clickable Table of Contents, cover page (WHO logo + Dian K Pro credit),
headers/footers with page numbers, acronyms, styled tables. AI-generated content fills
the relevant chapters; country-context sections are marked to be completed by the team."""
from __future__ import annotations
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config.settings import INSTITUTION_DARK, INSTITUTION_PRIMARY
from core.models import NISStrategy
from core.epi_components import EPI_COMPONENTS
from core.branding import logo_path, dk_logo_path, dk_credit

_DARK = RGBColor.from_string(INSTITUTION_DARK.lstrip("#"))
_PRIMARY = RGBColor.from_string(INSTITUTION_PRIMARY.lstrip("#"))
_GOLD = RGBColor.from_string("8A6D1B")
_HEX_DARK = INSTITUTION_DARK.lstrip("#")


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _field(paragraph, instr, placeholder="", dirty=False):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    if dirty:
        b.set(qn("w:dirty"), "true")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (b, it, sep, t, end):
        run._r.append(el)


def _enable_update_fields(doc):
    el = doc.settings.element
    for ex in el.findall(qn("w:updateFields")):
        el.remove(ex)
    upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
    el.append(upd)   # standard placement so Word honours it


def build_word(s: NISStrategy) -> bytes:
    lang = s.profile.language
    fr = lang == "fr"
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for _sec in doc.sections:   # 1-inch margins on all four sides
        _sec.top_margin = _sec.bottom_margin = _sec.left_margin = _sec.right_margin = Inches(1)
    period = f"{s.profile.nis_start_year}–{s.profile.nis_start_year + s.profile.nis_duration_years - 1}"
    years = s.profile.years
    ph = "À compléter par l’équipe pays" if fr else "To be completed by the country team"

    def H(text, level):
        # Numbered top-level chapters start on a new page -> distinct TOC page numbers.
        if level == 1 and str(text)[:1].isdigit():
            doc.add_page_break()
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = _DARK if level <= 1 else _PRIMARY
        return h

    def P(text=""):
        return doc.add_paragraph(text if text else ph)

    def field_val(label, val):
        p = doc.add_paragraph(); p.add_run(f"{label} : ").bold = True; p.add_run(str(val) if val else ph)

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = ""
            run = c.paragraphs[0].add_run(h); run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF"); run.font.size = Pt(9)
            _shade(c, _HEX_DARK)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                r = cells[i].paragraphs[0].add_run("" if v is None else str(v)); r.font.size = Pt(9)
        if widths:
            for rr in t.rows:
                for i, w in enumerate(widths):
                    rr.cells[i].width = Inches(w)
        return t

    # ---- footer: country/period + page number ----
    fp = doc.sections[0].footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run((f"SNV {s.profile.country_name} · {period}" if fr else
                f"NIS {s.profile.country_name} · {period}") + "    ·    ").font.size = Pt(8)
    _field(fp, "PAGE", "1")

    # ===================== COVER =====================
    lp = logo_path(lang)
    if lp:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(lp), width=Inches(2.7))
        except Exception:
            pass
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Stratégie Nationale de Vaccination" if fr else "National Immunization Strategy")
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = _DARK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{s.profile.country_name} · {period}"); r.font.size = Pt(18); r.font.color.rgb = _PRIMARY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\n{s.profile.ministry_name}\n{s.profile.epi_programme_name}\n"
              f"{'Date' if fr else 'Date'} : {s.profile.generation_date}").font.size = Pt(12)
    # Dian K Pro credit + logo
    for _ in range(2):
        doc.add_paragraph()
    dk = dk_logo_path()
    if dk:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(dk), width=Inches(0.7))
        except Exception:
            pass
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(dk_credit(lang)); r.font.size = Pt(10); r.italic = True; r.font.color.rgb = _GOLD
    doc.add_page_break()

    # ===================== TOC =====================
    # Title is NOT a Heading style, so the TOC does not list itself.
    ttl = doc.add_paragraph()
    r = ttl.add_run("Table des matières" if fr else "Table of contents")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = _DARK
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-3" \\h \\z \\u',
           "La table des matières se met à jour à l’ouverture (sinon : Ctrl+A puis F9)."
           if fr else "The table of contents updates on open (else: Ctrl+A then F9).", dirty=True)
    doc.add_page_break()

    # ===================== ABBREVIATIONS + EXEC SUMMARY =====================
    H("Liste des abréviations" if fr else "List of abbreviations", 1)
    acro = ([("PEV", "Programme Élargi de Vaccination"), ("SNV", "Stratégie Nationale de Vaccination"),
             ("FFOM", "Forces, Faiblesses, Opportunités, Menaces"), ("S&E", "Suivi et Évaluation"),
             ("IA2030", "Immunization Agenda 2030"), ("MAPI", "Manifestation Post-Immunisation"),
             ("MEV", "Maladie Évitable par la Vaccination"), ("SSP", "Soins de Santé Primaires"),
             ("OPS", "Objectif Prioritaire Stratégique")] if fr else
            [("EPI", "Expanded Programme on Immunization"), ("NIS", "National Immunization Strategy"),
             ("SWOT", "Strengths, Weaknesses, Opportunities, Threats"), ("M&E", "Monitoring & Evaluation"),
             ("IA2030", "Immunization Agenda 2030"), ("AEFI", "Adverse Event Following Immunization"),
             ("VPD", "Vaccine-Preventable Disease"), ("PHC", "Primary Health Care")])
    table(["Abréviation" if fr else "Abbreviation", "Signification" if fr else "Meaning"], acro, [1.6, 4.6])

    H("Résumé exécutif" if fr else "Executive summary", 1)
    P(s.vision.overall_objective)
    doc.add_paragraph(f"{'Objectifs stratégiques' if fr else 'Strategic objectives'} : {len(s.objectives)} · "
                      f"{'Interventions' if fr else 'Interventions'} : {len(s.interventions)} · "
                      f"{'Indicateurs' if fr else 'Indicators'} : {len(s.indicators)} · "
                      f"{'Activités' if fr else 'Activities'} : {len(s.activities)}")
    doc.add_page_break()

    # ===================== 1. INTRODUCTION =====================
    H("1. Introduction", 1)
    for sub in (["1.1 Situation géopolitique et démographique du pays",
                 "1.2 Situation socioéconomique du pays", "1.3 Organisation du système de santé",
                 "1.4 Principaux indicateurs de santé (mère et enfant)",
                 "1.5 Situation sécuritaire, populations déplacées et vulnérables",
                 "1.6 Parties prenantes"] if fr else
                ["1.1 Geopolitical and demographic situation", "1.2 Socioeconomic situation",
                 "1.3 Health system organization", "1.4 Key health indicators (maternal & child)",
                 "1.5 Security situation, displaced and vulnerable populations", "1.6 Stakeholders"]):
        H(sub, 2); P()

    # ===================== 2. EPI =====================
    H("2. Programme élargi de vaccination" if fr else "2. Expanded Programme on Immunization", 1)
    for sub in (["2.1 Historique", "2.2 Objectifs du programme", "2.3 Organigramme",
                 "2.4 Calendrier vaccinal", "2.5 Réalisations"] if fr else
                ["2.1 History", "2.2 Programme objectives", "2.3 Organogram",
                 "2.4 Immunization schedule", "2.5 Achievements"]):
        H(sub, 2); P()

    # ===================== 3. NIS PROCESS =====================
    H("3. Processus d’élaboration de la SNV" if fr else "3. NIS development process", 1)
    H("3.1 Équipe de développement et termes de référence" if fr else "3.1 Development team and ToR", 2); P()
    H("3.2 Processus d’élaboration de la SNV" if fr else "3.2 NIS development process", 2)
    P("Analyse situationnelle du PEV (7 composantes, 26 sous-composantes), analyse FFOM, analyse des causes "
      "profondes (méthode des POURQUOI), formulation des obstacles, théorie du changement, objectifs SMART, "
      "priorisation multicritère et cadre de S&E, en alignement avec l’IA2030 et la stratégie Gavi 6.0 (2026-2030)."
      if fr else
      "EPI situation analysis (7 components, 26 subcomponents), SWOT, root-cause analysis (5-Whys), obstacle "
      "formulation, theory of change, SMART objectives, multi-criteria prioritization and M&E, aligned with "
      "IA2030 and the Gavi 6.0 strategy (2026-2030).")
    H("3.3 Rôles des partenaires et parties prenantes" if fr else "3.3 Roles of partners and stakeholders", 2); P()

    # ===================== 4. SITUATION ANALYSIS =====================
    H("4. Analyse de la situation" if fr else "4. Situation analysis", 1)
    H("4.1 Analyse du secteur de la santé" if fr else "4.1 National health sector analysis", 2)
    for sub in (["4.1.1 Gouvernance et leadership", "4.1.2 Personnel de santé",
                 "4.1.3 Prestation des services de santé", "4.1.4 Système d’information sanitaire",
                 "4.1.5 Financement de la santé"] if fr else
                ["4.1.1 Governance and leadership", "4.1.2 Health workforce",
                 "4.1.3 Health service delivery", "4.1.4 Health management information system",
                 "4.1.5 Health financing"]):
        H(sub, 3); P()
    H("4.2 Analyse des composantes du PEV (FFOM)" if fr else "4.2 EPI components analysis (SWOT)", 2)
    sw_by = {x.subcomponent_code: x for x in s.swot}
    rc_by = {}
    for rc in s.root_causes:
        rc_by.setdefault(rc.subcomponent_code, []).append(rc)
    quad = [("Forces" if fr else "Strengths", "strengths"),
            ("Faiblesses" if fr else "Weaknesses", "weaknesses"),
            ("Opportunités" if fr else "Opportunities", "opportunities"),
            ("Menaces" if fr else "Threats", "threats")]
    ci = 0
    for comp in EPI_COMPONENTS:
        ci += 1
        subs_with = [sub for sub in comp.subcomponents
                     if (it := sw_by.get(sub.code)) and any([it.strengths, it.weaknesses,
                                                             it.opportunities, it.threats])]
        if not subs_with:
            continue
        H(f"4.2.{ci} {comp.label(lang)}", 3)
        for sub in subs_with:
            it = sw_by[sub.code]
            H(sub.label(lang), 4)
            for label, attr in quad:
                vals = getattr(it, attr)
                if vals:
                    p = doc.add_paragraph(); p.add_run(f"{label} : ").bold = True; p.add_run("; ".join(vals))
            for rc in rc_by.get(sub.code, []):
                p = doc.add_paragraph()
                p.add_run(("Cause profonde : " if fr else "Root cause: ")).bold = True
                p.add_run(f"{' → '.join(rc.whys)} ⇒ {rc.final_why}")

    # ===================== 5. VISION =====================
    H("5. Vision, objectif global et objectif général de la SNV" if fr else
      "5. NIS vision, global goal and overall objective", 1)
    H("5.1 Vision, objectif global et objectif général" if fr else
      "5.1 Vision, global goal and overall objective", 2)
    field_val("Vision", s.vision.vision)
    field_val("But de la SNV" if fr else "NIS goal", s.vision.goal)
    field_val("Objectif général" if fr else "Overall objective", s.vision.overall_objective)
    for sub in (["5.2 Vaccination tout au long de la vie", "5.3 Vaccination contre la COVID-19",
                 "5.4 Introduction de nouveaux vaccins et cibles de couverture",
                 "5.5 Vaccination en situations de conflit et d’urgence"] if fr else
                ["5.2 Life-course vaccination", "5.3 COVID-19 vaccination",
                 "5.4 New vaccine introduction and coverage targets",
                 "5.5 Vaccination in conflict and emergency settings"]):
        H(sub, 2); P()

    # ===================== 6. OBJECTIVES & INTERVENTIONS =====================
    H("6. Objectifs stratégiques prioritaires et interventions principales" if fr else
      "6. Strategic priority objectives and main interventions", 1)
    H("6.1 Objectifs prioritaires stratégiques et liens avec l’IA2030" if fr else
      "6.1 Strategic priority objectives and links to IA2030", 2)
    for o in s.objectives:
        H(o.objective_text or o.obj_id, 3)
        field_val("Obstacle principal" if fr else "Main obstacle", o.main_obstacle)
        field_val("Résultat visionnaire" if fr else "Visionary result", o.visionary_result)
    if not s.objectives:
        P()
    H("6.2 Interventions principales par composante" if fr else "6.2 Main interventions by component", 2)
    iv_by_obj = {}
    for iv in s.interventions:
        iv_by_obj.setdefault(iv.objective_id, []).append(iv)
    for o in s.objectives:
        ivs = iv_by_obj.get(o.obj_id, [])
        if not ivs:
            continue
        H(o.objective_text or o.obj_id, 3)
        for iv in ivs:
            p = doc.add_paragraph()
            p.add_run(f"{iv.title}  [{getattr(iv.priority_level, 'value', iv.priority_level)}]").bold = True
            for label, val in [("Justification" if fr else "Rationale", iv.rationale),
                               ("Impact attendu" if fr else "Expected impact", iv.expected_impact),
                               ("Prérequis" if fr else "Prerequisites", "; ".join(iv.prerequisites)),
                               ("Risques" if fr else "Risks", "; ".join(iv.risks)),
                               ("Partenaires" if fr else "Partners", "; ".join(iv.partners)),
                               ("Calendrier" if fr else "Timeline",
                                ", ".join(str(years[i]) for i in range(len(years)) if iv.timeline.get(f"Y{i+1}")))]:
                if val:
                    pp = doc.add_paragraph(style="List Bullet")
                    pp.add_run(f"{label} : ").bold = True; pp.add_run(str(val))
            refs = "; ".join(f"{e.document_name} ({e.locator})" for e in iv.evidence if e.document_name)
            if refs:
                pp = doc.add_paragraph(style="List Bullet")
                pp.add_run(("Références : " if fr else "References: ")).bold = True
                pp.add_run(refs).italic = True
    if not s.interventions:
        P()

    # ===================== 7. M&E =====================
    H("7. Cadre de suivi et d’évaluation pour l’action" if fr else "7. Monitoring & evaluation framework", 1)
    H("7.1 Indicateurs clés par objectif" if fr else "7.1 Key indicators by objective", 2)
    if s.indicators:
        cols = ["Indicateur" if fr else "Indicator", "Type", "Base" if fr else "Baseline"] + [str(y) for y in years]
        rows = [[ind.name, str(ind.indicator_type), ind.baseline]
                + [ind.targets.get(f"Y{i+1}", "") for i in range(len(years))] for ind in s.indicators]
        table(cols, rows)
    else:
        P()
    H("7.2 Tableau de bord et mécanismes de retour d’information" if fr else
      "7.2 Country dashboard and feedback mechanisms", 2); P()
    H("7.3 Notification et engagement régional et mondial" if fr else
      "7.3 Regional and global reporting and engagement", 2); P()

    # ===================== 8. COSTS =====================
    H("8. Coûts et financement de la SNV" if fr else "8. NIS costs and financing", 1)
    for sub in (["8.1 Coûts de la SNV par composante et par année", "8.2 Analyse des écarts de coûts",
                 "8.3 Sources de financement par composante et par année", "8.4 Durabilité financière"] if fr else
                ["8.1 NIS costs by component and year", "8.2 Costing gap analysis",
                 "8.3 Funding sources by component and year", "8.4 Financial sustainability"]):
        H(sub, 2)
        P("À réaliser avec l’outil de chiffrage NIS.COST." if fr else "To be completed with the NIS.COST tool.")

    # ===================== 9. ANNEXES + YEAR-1 ACTIVITIES =====================
    H("9. Annexes" if fr else "9. Annexes", 1)
    H("Sources utilisées" if fr else "Sources used", 2)
    for d in s.documents:
        doc.add_paragraph(f"{d.name} ({d.doc_category})", style="List Bullet")
    if not s.documents:
        P()
    H("Chronogramme annuel des activités du PEV" if fr else "Annual EPI activity timeline", 2)
    if s.activities:
        rows = []
        for a in s.activities:
            yrs = ", ".join(str(years[i]) for i in range(len(years)) if a.years.get(f"Y{i+1}"))
            rows.append([a.activity, a.implementation_level, a.lead, yrs])
        table(["Activité clé" if fr else "Key activity", "Niveau" if fr else "Level",
               "Responsable" if fr else "Lead", "Années" if fr else "Years"], rows, [3.0, 1.6, 1.4, 1.0])
    else:
        P()

    _enable_update_fields(doc)
    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()
