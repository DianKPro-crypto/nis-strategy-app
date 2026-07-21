"""Step 11/12 exports:
- build_narrative_word: full AI-written NIS narrative report (prose per section + key tables).
- build_financial_word: financial report generated from NIS.COST data.
- build_qa_word: AI quality-assurance report, with issues to fix marked in RED.
All premium, with cover (WHO logo + Dian K Pro credit), auto clickable TOC and page numbers.
"""
from __future__ import annotations
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config.settings import INSTITUTION_DARK, INSTITUTION_PRIMARY
from core.models import NISStrategy
from core.branding import logo_path, dk_logo_path, dk_credit
from core.ai_engine import NARRATIVE_SECTIONS

_DARK = RGBColor.from_string(INSTITUTION_DARK.lstrip("#"))
_PRIMARY = RGBColor.from_string(INSTITUTION_PRIMARY.lstrip("#"))
_GOLD = RGBColor.from_string("8A6D1B")
_RED = RGBColor.from_string("C00000")
_HEX_DARK = INSTITUTION_DARK.lstrip("#")


def _field(p, instr, placeholder="", dirty=False):
    run = p.add_run()
    for kind, txt in (("begin", None), ("instr", instr), ("separate", None), ("text", placeholder), ("end", None)):
        if kind == "instr":
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = txt
        elif kind == "text":
            el = OxmlElement("w:t"); el.text = txt
        else:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), kind)
            if kind == "begin" and dirty:      # forces Word to recompute the field on open
                el.set(qn("w:dirty"), "true")
        run._r.append(el)


def _update_fields(doc):
    # Standard placement (append) so Word honours "recompute fields on open" (schema-valid).
    el = doc.settings.element
    for ex in el.findall(qn("w:updateFields")):
        el.remove(ex)
    upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
    el.append(upd)


def _footer(section, txt):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(txt + "    ·    ").font.size = Pt(8)
    _field(p, "PAGE", "1")


def _cover(doc, s, title):
    fr = s.profile.language == "fr"
    lp = logo_path(s.profile.language)
    if lp:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(lp), width=Inches(2.7))
        except Exception:
            pass
    for _ in range(3):
        doc.add_paragraph()
    period = f"{s.profile.nis_start_year}–{s.profile.nis_start_year + s.profile.nis_duration_years - 1}"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = _DARK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{s.profile.country_name} · {period}"); r.font.size = Pt(16); r.font.color.rgb = _PRIMARY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\n{s.profile.ministry_name}\n{s.profile.epi_programme_name}\n{s.profile.generation_date}").font.size = Pt(12)
    for _ in range(2):
        doc.add_paragraph()
    dk = dk_logo_path()
    if dk:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(dk), width=Inches(0.7))
        except Exception:
            pass
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(dk_credit(s.profile.language)); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = _GOLD
    doc.add_page_break()


def _H(doc, text, level, numbered_break=True):
    if level == 1 and numbered_break:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = _DARK if level <= 1 else _PRIMARY
    return h


def _toc(doc, fr):
    _H(doc, "Table des matières" if fr else "Table of contents", 1, numbered_break=False)
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-2" \\h \\z \\u',
           "Se met à jour à l’ouverture (sinon : Ctrl+A puis F9)." if fr
           else "Updates on open (else: Ctrl+A then F9).", dirty=True)
    doc.add_page_break()


def _add_rich(paragraph, text):
    """Add text to a paragraph, rendering **bold** markdown as real bold runs."""
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part.replace("**", ""))


def _clean_head(t):
    return t.strip().lstrip("#").strip().replace("**", "").strip("* ")


def _prose(doc, text):
    """Render AI prose: '## '/'### ' -> Heading 2/3, '- ' -> bullets, **bold** -> bold runs."""
    for raw in (text or "").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#### ") or line.startswith("### "):
            _H(doc, _clean_head(line), 3, numbered_break=False)
        elif line.startswith("## ") or line.startswith("# "):
            _H(doc, _clean_head(line), 2, numbered_break=False)
        elif line[:2] in ("- ", "* ") or line.startswith("• "):
            _add_rich(doc.add_paragraph(style="List Bullet"), line[2:].strip())
        else:
            _add_rich(doc.add_paragraph(), line)


def _new_doc(s):
    doc = Document()
    # --- Premium typography (UN-publication feel: clean, justified body, colour-coded headings) ---
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1A1A1A")
    pf = normal.paragraph_format
    pf.line_spacing = 1.15; pf.space_after = Pt(6); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color in (("Heading 1", 17, _DARK), ("Heading 2", 13, _PRIMARY), ("Heading 3", 11.5, _DARK)):
        try:
            stl = doc.styles[name]
            stl.font.name = "Calibri"; stl.font.size = Pt(size); stl.font.bold = True; stl.font.color.rgb = color
            hp = stl.paragraph_format
            hp.space_before = Pt(14 if name == "Heading 1" else 9); hp.space_after = Pt(4); hp.keep_with_next = True
        except Exception:
            pass
    for sec in doc.sections:   # 1-inch margins on all four sides
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    fr = s.profile.language == "fr"
    period = f"{s.profile.nis_start_year}–{s.profile.nis_start_year + s.profile.nis_duration_years - 1}"
    _footer(doc.sections[0], (f"SNV {s.profile.country_name} · {period}" if fr
                              else f"NIS {s.profile.country_name} · {period}"))
    return doc


def _set_fill(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexfill); tcPr.append(shd)


def _mini_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        pr = c.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = pr.add_run(h); r.bold = True; r.font.color.rgb = RGBColor.from_string("FFFFFF"); r.font.size = Pt(9)
        _set_fill(c, _HEX_DARK)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pr.add_run("" if v is None else str(v)).font.size = Pt(9)
            if ri % 2 == 0:                     # zebra striping for readability
                _set_fill(cells[i], "EEF3F8")


def _seq_caption(doc, label, title):
    """Numbered caption ('Tableau 1. …' / 'Figure 1. …') via a SEQ field, so Word can build
    the lists of tables/figures automatically."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Caption"]
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label} "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _PRIMARY
    _field(p, f"SEQ {label} \\* ARABIC", "1")
    r2 = p.add_run(f". {title}"); r2.font.size = Pt(9); r2.italic = True
    return p


def _list_of(doc, heading, label):
    _H(doc, heading, 2, numbered_break=False)
    p = doc.add_paragraph()
    _field(p, f'TOC \\h \\z \\c "{label}"', "(se met à jour à l’ouverture)", dirty=True)


def _titled_table(doc, title, headers, rows):
    _seq_caption(doc, "Tableau", title)
    _mini_table(doc, headers, rows)


# --------------------------------------------------------------------------- #
def build_narrative_word(s: NISStrategy) -> bytes:
    fr = s.profile.language == "fr"
    years = s.profile.years
    doc = _new_doc(s)
    _cover(doc, s, "Stratégie Nationale de Vaccination" if fr else "National Immunization Strategy")
    _toc(doc, fr)
    # ---- Front matter: abbreviations/glossary + lists of tables & figures ----
    _H(doc, "Sigles et abréviations" if fr else "Abbreviations", 1)
    _mini_table(doc, ["Sigle" if fr else "Acronym", "Signification" if fr else "Meaning"],
                [("PEV", "Programme Élargi de Vaccination"), ("SNV", "Stratégie Nationale de Vaccination"),
                 ("FFOM", "Forces, Faiblesses, Opportunités, Menaces"), ("S&E", "Suivi et Évaluation"),
                 ("IA2030", "Immunization Agenda 2030"), ("Gavi", "Alliance mondiale pour les vaccins"),
                 ("MAPI", "Manifestation Post-Immunisation"), ("MEV", "Maladie Évitable par la Vaccination"),
                 ("SSP", "Soins de Santé Primaires"), ("OPS", "Objectif Prioritaire Stratégique"),
                 ("POA", "Plan Opérationnel Annuel"), ("RSS", "Renforcement du Système de Santé")]
                if fr else
                [("EPI", "Expanded Programme on Immunization"), ("NIS", "National Immunization Strategy"),
                 ("SWOT", "Strengths, Weaknesses, Opportunities, Threats"), ("M&E", "Monitoring & Evaluation"),
                 ("IA2030", "Immunization Agenda 2030"), ("Gavi", "Global Alliance for Vaccines"),
                 ("AEFI", "Adverse Event Following Immunization"), ("VPD", "Vaccine-Preventable Disease"),
                 ("PHC", "Primary Health Care"), ("AOP", "Annual Operational Plan")])
    _list_of(doc, "Liste des tableaux" if fr else "List of tables", "Tableau")
    doc.add_page_break()
    from core.epi_components import EPI_COMPONENTS, find_subcomponent
    sw_by = {x.subcomponent_code: x for x in s.swot}

    for n, (key, tfr, ten) in enumerate(NARRATIVE_SECTIONS, 1):
        _H(doc, f"{n}. {tfr if fr else ten}", 1)
        _prose(doc, s.narrative.get(key, "") or ("À compléter par l’équipe pays" if fr
                                                  else "To be completed by the country team"))
        # ---- section tables & figures ----
        if key == "situation" and s.swot:
            p = doc.add_paragraph()
            r = p.add_run("La synthèse FFOM complète par sous-composante figure en Annexe A." if fr
                          else "The full SWOT summary by subcomponent is provided in Annex A.")
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = _PRIMARY
        if key == "objectives" and s.objectives:
            _titled_table(doc, "Objectifs stratégiques prioritaires" if fr else "Priority strategic objectives",
                          ["ID", "Objectif" if fr else "Objective", "Obstacle principal" if fr else "Main obstacle"],
                          [[o.obj_id, o.objective_text, o.main_obstacle] for o in s.objectives])
        if key == "interventions" and s.interventions:
            _titled_table(doc, "Interventions prioritaires et calendrier" if fr else "Priority interventions",
                          ["Intervention", "Priorité" if fr else "Priority", "Calendrier" if fr else "Timeline"],
                          [[iv.title, getattr(iv.priority_level, "value", iv.priority_level),
                            ", ".join(str(years[k]) for k in range(len(years)) if iv.timeline.get(f"Y{k+1}"))]
                           for iv in s.interventions])
        if key == "me" and s.indicators:
            cols = ["Indicateur" if fr else "Indicator", "Base" if fr else "Baseline"] + [str(y) for y in years]
            _titled_table(doc, "Indicateurs de suivi-évaluation et cibles" if fr else "M&E indicators and targets",
                          cols, [[i.name, i.baseline] + [i.targets.get(f"Y{k+1}", "")
                                 for k in range(len(years))] for i in s.indicators])

    if s.financial_report:
        _H(doc, f"{len(NARRATIVE_SECTIONS)+1}. " + ("Rapport financier" if fr else "Financial report"), 1)
        _prose(doc, s.financial_report)

    # ===================== ANNEXES =====================
    _H(doc, ("Annexes" if fr else "Annexes"), 1)
    _H(doc, "Annexe A — Synthèse FFOM par sous-composante" if fr else "Annex A — SWOT summary by subcomponent", 2)
    if s.swot:
        def _cell(items):  # all items, joined; capped to avoid an oversized cell
            txt = " ; ".join(x for x in items if x)
            return (txt[:600] + "…") if len(txt) > 600 else (txt or "—")
        rows = []
        for c in EPI_COMPONENTS:
            for sub in c.subcomponents:   # ALL 26 subcomponents
                it = sw_by.get(sub.code)
                rows.append([sub.label(s.profile.language),
                             _cell(it.strengths) if it else "—", _cell(it.weaknesses) if it else "—",
                             _cell(it.opportunities) if it else "—", _cell(it.threats) if it else "—"])
        _titled_table(doc, "Synthèse FFOM par sous-composante" if fr else "SWOT summary by subcomponent",
                      ["Sous-composante" if fr else "Subcomponent", "Forces", "Faiblesses",
                       "Opportunités", "Menaces"], rows)
    else:
        doc.add_paragraph("À compléter." if fr else "To be completed.")
    _H(doc, "Annexe B — Chronogramme des activités" if fr else "Annex B — Activity timeline", 2)
    if s.activities:
        _titled_table(doc, "Chronogramme des activités" if fr else "Activity timeline",
                      ["Activité" if fr else "Activity", "Niveau" if fr else "Level",
                       "Responsable" if fr else "Lead", "Années" if fr else "Years"],
                      [[a.activity, a.implementation_level, a.lead,
                        ", ".join(str(years[k]) for k in range(len(years)) if a.years.get(f"Y{k+1}"))]
                       for a in s.activities])
    else:
        doc.add_paragraph("À compléter." if fr else "To be completed.")
    _H(doc, "Annexe C — Cadre de S&E détaillé" if fr else "Annex C — Detailed M&E framework", 2)
    if s.indicators:
        _titled_table(doc, "Cadre de S&E détaillé" if fr else "Detailed M&E framework",
                      ["Indicateur" if fr else "Indicator", "Type", "Définition" if fr else "Definition",
                       "Source" if fr else "Source", "Fréquence" if fr else "Frequency"],
                      [[i.name, str(i.indicator_type), i.definition, i.data_source, i.frequency]
                       for i in s.indicators])
    else:
        doc.add_paragraph("À compléter." if fr else "To be completed.")
    _H(doc, "Annexe D — Analyse des causes profondes" if fr else "Annex D — Root-cause analysis", 2)
    if s.root_causes:
        _titled_table(doc, "Analyse des causes profondes" if fr else "Root-cause analysis",
                      ["Sous-comp." if fr else "Sub-comp.", "Faiblesse" if fr else "Weakness",
                       "POURQUOI" if fr else "WHYs", "Dernier POURQUOI" if fr else "Last WHY",
                       "Problème principal" if fr else "Main problem"],
                      [[rc.subcomponent_code, rc.weakness, " → ".join(rc.whys), rc.final_why,
                        rc.main_problem] for rc in s.root_causes[:40]])
    else:
        doc.add_paragraph("À compléter." if fr else "To be completed.")
    _H(doc, "Annexe E — Sources et références" if fr else "Annex E — Sources and references", 2)
    for d in s.documents:
        doc.add_paragraph(f"{d.name} ({d.doc_category})", style="List Bullet")
    if not s.documents:
        doc.add_paragraph("À compléter." if fr else "To be completed.")

    _update_fields(doc)
    buf = BytesIO(); doc.save(buf); return buf.getvalue()


def build_financial_word(s: NISStrategy) -> bytes:
    fr = s.profile.language == "fr"
    doc = _new_doc(s)
    _cover(doc, s, "Rapport financier de la SNV (NIS.COST)" if fr else "NIS financial report (NIS.COST)")
    _H(doc, "Rapport financier" if fr else "Financial report", 1)
    _prose(doc, s.financial_report or ("À générer à partir du fichier NIS.COST." if fr
                                       else "To be generated from the NIS.COST file."))
    _update_fields(doc); buf = BytesIO(); doc.save(buf); return buf.getvalue()


def build_qa_word(s: NISStrategy, qa: dict) -> bytes:
    """QA report: issues to add/improve are marked in RED."""
    fr = s.profile.language == "fr"
    doc = _new_doc(s)
    _cover(doc, s, "Rapport d’assurance qualité de la SNV" if fr else "NIS quality-assurance report")
    _H(doc, ("Synthèse et score" if fr else "Summary and score"), 1)
    p = doc.add_paragraph()
    r = p.add_run(f"{'Score global' if fr else 'Overall score'} : {qa.get('score', '—')}/100")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = _DARK
    doc.add_paragraph(qa.get("overall", ""))
    _H(doc, ("Points à ajouter / améliorer (en rouge)" if fr else "Items to add / improve (in red)"), 1)
    findings = qa.get("findings", []) or []
    order = {"critique": 0, "critical": 0, "majeur": 1, "major": 1, "mineur": 2, "minor": 2}
    for f in sorted(findings, key=lambda x: order.get(str(x.get("severity", "")).lower(), 3)):
        sev = str(f.get("severity", "")).lower()
        red = sev in ("critique", "critical", "majeur", "major")
        p = doc.add_paragraph(style="List Bullet")
        head = p.add_run(f"[{f.get('severity', '')}] {f.get('section', '')} — {f.get('issue', '')}")
        head.bold = True
        if red:
            head.font.color.rgb = _RED
        rec = doc.add_paragraph()
        rr = rec.add_run(f"→ {('Recommandation' if fr else 'Recommendation')} : {f.get('recommendation', '')}")
        if red:
            rr.font.color.rgb = _RED
    if not findings:
        doc.add_paragraph("Aucun problème majeur détecté." if fr else "No major issue detected.")
    _update_fields(doc); buf = BytesIO(); doc.save(buf); return buf.getvalue()


def build_narrative_pdf(s: NISStrategy) -> bytes:
    """SNV narrative as a PDF with REAL page numbers + clickable TOC (ReportLab computes pagination —
    no dependency on Word field updates)."""
    import re as _re
    from core.epi_components import EPI_COMPONENTS
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer,
                                    Table, TableStyle, PageBreak)
    from reportlab.platypus.tableofcontents import TableOfContents

    fr = s.profile.language == "fr"
    DARK = colors.HexColor("#003366"); PRIM = colors.HexColor("#0093D5")
    period = f"{s.profile.nis_start_year}–{s.profile.nis_start_year + s.profile.nis_duration_years - 1}"
    header = (f"SNV {s.profile.country_name} {period}" if fr else f"NIS {s.profile.country_name} {period}")
    lp = logo_path(s.profile.language)
    years = s.profile.years
    st = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=st["Heading1"], textColor=DARK, spaceBefore=10, spaceAfter=6, fontSize=15)
    h2 = ParagraphStyle("h2", parent=st["Heading2"], textColor=PRIM, spaceBefore=8, spaceAfter=4, fontSize=12)
    h3 = ParagraphStyle("h3", parent=st["Heading3"], textColor=DARK, fontSize=10.5, spaceBefore=5, spaceAfter=2)
    body = ParagraphStyle("body", parent=st["BodyText"], fontSize=9.7, leading=13.5, spaceAfter=5, alignment=4)
    cell = ParagraphStyle("cell", parent=st["BodyText"], fontSize=8, leading=10, spaceAfter=0)
    cellh = ParagraphStyle("cellh", parent=cell, textColor=colors.white, fontName="Helvetica-Bold")

    def esc(v):
        return (str(v if v is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def rich(text):
        return _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", esc(text))

    class _DocT(BaseDocTemplate):
        def afterFlowable(self, fl):
            lvl = getattr(fl, "_toc_level", None)
            if lvl is not None:
                key = f"h{id(fl)}"; self.canv.bookmarkPage(key)
                self.notify("TOCEntry", (lvl, fl.getPlainText(), self.page, key))
                try: self.canv.addOutlineEntry(fl.getPlainText(), key, level=lvl, closed=False)
                except Exception: pass

    def _decorate(cv, d):
        cv.saveState(); cv.setFont("Helvetica", 8); cv.setFillColor(DARK)
        if lp:
            try:
                cv.drawImage(str(lp), 2 * cm, A4[1] - 1.9 * cm, width=3.2 * cm, height=0.85 * cm,
                             preserveAspectRatio=True, mask="auto")
            except Exception: pass
        cv.drawRightString(A4[0] - 2 * cm, A4[1] - 1.3 * cm, header)
        cv.setStrokeColor(PRIM); cv.line(2 * cm, A4[1] - 2.0 * cm, A4[0] - 2 * cm, A4[1] - 2.0 * cm)
        cv.setFillColor(colors.grey); cv.drawCentredString(A4[0] / 2, 1 * cm, f"{d.page}")
        cv.restoreState()

    buf = BytesIO()
    doc = _DocT(buf, pagesize=A4, title=header)
    frame = Frame(2 * cm, 2 * cm, A4[0] - 4 * cm, A4[1] - 2.6 * cm - 2 * cm, id="body")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=_decorate)])
    story = []

    def chapter(t):
        p = Paragraph(esc(t), h1); p._toc_level = 0; story.append(p)

    def prose(text):
        for raw in (text or "").split("\n"):
            line = raw.strip()
            if not line:
                continue
            if line.startswith("### ") or line.startswith("#### "):
                story.append(Paragraph(rich(line.lstrip("#").strip().strip("*")), h3))
            elif line.startswith("## ") or line.startswith("# "):
                p = Paragraph(rich(line.lstrip("#").strip().strip("*")), h2); p._toc_level = 1; story.append(p)
            elif line[:2] in ("- ", "* ") or line.startswith("• "):
                story.append(Paragraph("•&nbsp; " + rich(line[2:].strip()), body))
            else:
                story.append(Paragraph(rich(line), body))

    def _cellval(v):   # bound cell content so no single row can exceed a page (avoids LayoutError)
        if isinstance(v, (list, tuple)):
            items = [esc(str(x)[:130]) for x in v if x][:4]
            return ("• " + "<br/>• ".join(items)) if items else "—"
        return esc(str(v)[:350]) or "—"

    def table(headers, rows, widths):
        data = [[Paragraph(esc(h), cellh) for h in headers]]
        for row in rows:
            data.append([Paragraph(_cellval(v), cell) for v in row])
        t = Table(data, colWidths=widths, repeatRows=1, splitByRow=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), DARK), ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B9C6D4")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF3F8")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
        story.append(Spacer(1, 4)); story.append(t); story.append(Spacer(1, 6))

    W = A4[0] - 4 * cm
    # ---- Cover ----
    story += [Spacer(1, 5 * cm),
              Paragraph("Stratégie Nationale de Vaccination" if fr else "National Immunization Strategy",
                        ParagraphStyle("t", parent=h1, fontSize=24, alignment=1, textColor=DARK)),
              Spacer(1, .4 * cm),
              Paragraph(f"{esc(s.profile.country_name)} · {period}",
                        ParagraphStyle("t2", parent=body, fontSize=14, alignment=1)),
              Spacer(1, .3 * cm),
              Paragraph(f"{esc(s.profile.ministry_name)}<br/>{esc(s.profile.epi_programme_name)}"
                        f"<br/>{esc(s.profile.generation_date)}",
                        ParagraphStyle("t3", parent=body, alignment=1)),
              Spacer(1, 2.5 * cm),
              Paragraph(esc(dk_credit(s.profile.language)),
                        ParagraphStyle("dk", parent=body, alignment=1, fontSize=9,
                                       textColor=colors.HexColor("#8A6D1B"))),
              PageBreak()]
    # ---- Clickable TOC (real page numbers via multiBuild) ----
    story.append(Paragraph("Table des matières" if fr else "Table of contents", h1))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("toc0", fontSize=11, leading=18, textColor=DARK, leftIndent=6, firstLineIndent=-6),
        ParagraphStyle("toc1", fontSize=9.5, leading=14, textColor=PRIM, leftIndent=22, firstLineIndent=-6)]
    story += [toc, PageBreak()]

    # ---- Chapters ----
    for n, (key, tfr, ten) in enumerate(NARRATIVE_SECTIONS, 1):
        chapter(f"{n}. {tfr if fr else ten}")
        prose(s.narrative.get(key, "") or ("À compléter par l’équipe pays" if fr
                                           else "To be completed by the country team"))
        if key == "objectives" and s.objectives:
            table(["ID", "Objectif" if fr else "Objective", "Obstacle" if fr else "Obstacle"],
                  [[o.obj_id, o.objective_text, o.main_obstacle] for o in s.objectives],
                  [1.8 * cm, W - 6.3 * cm, 4.5 * cm])
        if key == "interventions" and s.interventions:
            table(["Intervention", "Prio." if fr else "Prio.", "Calendrier" if fr else "Timeline"],
                  [[iv.title, getattr(iv.priority_level, "value", iv.priority_level),
                    ", ".join(str(years[k]) for k in range(len(years)) if iv.timeline.get(f"Y{k+1}"))]
                   for iv in s.interventions], [W - 5.5 * cm, 1.8 * cm, 3.7 * cm])
        if key == "me" and s.indicators:
            cols = ["Indicateur" if fr else "Indicator", "Base"] + [str(y) for y in years]
            yw = (W - 5.5 * cm) / max(1, len(years))
            table(cols, [[i.name, i.baseline] + [i.targets.get(f"Y{k+1}", "") for k in range(len(years))]
                         for i in s.indicators], [3.7 * cm, 1.8 * cm] + [yw] * len(years))

    if s.financial_report:
        chapter(f"{len(NARRATIVE_SECTIONS)+1}. " + ("Rapport financier" if fr else "Financial report"))
        prose(s.financial_report)

    # ---- Annexes (detailed tables) ----
    story.append(PageBreak()); chapter("Annexes")
    if s.swot:
        story.append(Paragraph("Annexe A — Synthèse FFOM par sous-composante" if fr
                               else "Annex A — SWOT summary by subcomponent", h2))
        sw = {x.subcomponent_code: x for x in s.swot}
        rows = []
        for c in EPI_COMPONENTS:
            for sub in c.subcomponents:
                it = sw.get(sub.code)
                rows.append([sub.label(s.profile.language),
                             it.strengths if it else "—", it.weaknesses if it else "—",
                             it.opportunities if it else "—", it.threats if it else "—"])
        table(["Sous-comp." if fr else "Subcomp.", "Forces", "Faiblesses", "Opport.", "Menaces"],
              rows, [3.2 * cm, (W - 3.2 * cm) / 4] * 1 + [(W - 3.2 * cm) / 4] * 4)
    if s.activities:
        story.append(Paragraph("Annexe B — Chronogramme des activités" if fr
                               else "Annex B — Activity timeline", h2))
        table(["Activité" if fr else "Activity", "Niveau" if fr else "Level", "Resp.",
               "Années" if fr else "Years"],
              [[a.activity, a.implementation_level, a.lead,
                ", ".join(str(years[k]) for k in range(len(years)) if a.years.get(f"Y{k+1}"))]
               for a in s.activities], [W - 8.5 * cm, 3 * cm, 2.5 * cm, 3 * cm])
    if s.root_causes:
        story.append(Paragraph("Annexe C — Analyse des causes profondes" if fr
                               else "Annex C — Root-cause analysis", h2))
        table(["Sous-comp." if fr else "Sub-comp.", "Faiblesse" if fr else "Weakness",
               "POURQUOI" if fr else "WHYs", "Problème principal" if fr else "Main problem"],
              [[rc.subcomponent_code, rc.weakness, " → ".join(rc.whys),
                rc.main_problem or rc.final_why] for rc in s.root_causes[:40]],
              [2.2 * cm, (W - 6.6 * cm) / 2, (W - 6.6 * cm) / 2, 4.4 * cm])
    if s.documents:
        story.append(Paragraph("Annexe D — Sources et références" if fr else "Annex D — Sources", h2))
        for d in s.documents:
            story.append(Paragraph(f"•&nbsp; {esc(d.name)} ({esc(d.doc_category)})", body))

    doc.multiBuild(story)   # multiple passes -> correct TOC page numbers
    return buf.getvalue()
